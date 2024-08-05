from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.chrome.options import Options as ChromeOptions
from docx import Document
from bs4 import BeautifulSoup
import difflib
import time

def get_webpage_content(url):
    # Create a ChromeDriver instance
    driver = webdriver.Chrome()

    # Navigate to the webpage
    driver.get(url)

    # Wait for the webpage to load (you might need to adjust the wait time)
    time.sleep(5)

    # Extract the content from the webpage
    webpage_source = driver.page_source

    # Parse HTML with BeautifulSoup and extract only text content
    soup = BeautifulSoup(webpage_source, 'html.parser')
    webpage_text_content = soup.get_text(separator='\n')

    # Close the Chrome window
    driver.quit()

    return webpage_text_content

def get_docx_content(docx_path):
    # Read the content from the Word document
    doc = Document(docx_path)
    doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    return doc_content

def show_differences(content1, content2):
    # Use difflib to compare the two sets of content
    differ = difflib.Differ()
    diff = list(differ.compare(content1.splitlines(), content2.splitlines()))

    # Collect lines with differences
    webpage_only = [line[2:] for line in diff if line.startswith('+ ')]
    doc_only = [line[2:] for line in diff if line.startswith('- ')]

    return webpage_only, doc_only

def log_differences_to_docx(webpage_only, doc_only, output_path):
    # Create a new Word document
    doc = Document()

    # Add headings for additions and deletions
    doc.add_heading('Things Missing from Webpage', level=1)
    for line in webpage_only:
        doc.add_paragraph(line)

    doc.add_heading('Things Missing from Document', level=1)
    for line in doc_only:
        doc.add_paragraph(line)

    # Save the document
    doc.save(output_path)

def main():
    # URL of the webpage to test
    webpage_url = 'https://emploin.in/UnifiedSSO'

    # Path to the Word document
    docx_path = r"C:\Users\User1\Downloads\Web_Docs\UnifiedProfileSSOSolutionsPageContentv3.docx"

    # Get content from the webpage (text only, without HTML tags)
    webpage_content = get_webpage_content(webpage_url)

    # Get content from the Word document
    doc_content = get_docx_content(docx_path)

    # Show the differences
    webpage_only, doc_only = show_differences(webpage_content, doc_content)

    # Log differences to a new Word document
    log_differences_to_docx(webpage_only, doc_only, 'differences_output.docx')

if __name__ == "__main__":
    main()
