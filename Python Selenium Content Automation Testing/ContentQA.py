from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from docx import Document
from bs4 import BeautifulSoup
import difflib
import time
import pandas as pd

def get_webpage_content(url):
    # Create a ChromeDriver instance
    driver = webdriver.Chrome()

    # Navigate to the webpage
    driver.get(url)

    # Wait for the webpage to load (you might need to adjust the wait time)
    time.sleep(5)

    # Extract the content from the webpage
    webpage_source = driver.page_source

    # Parse HTML with BeautifulSoup and extract only text content
    soup = BeautifulSoup(webpage_source, 'html.parser')
    webpage_text_content = soup.get_text(separator='\n')

    # Close the Chrome window
    driver.quit()

    return webpage_text_content

def get_docx_content(docx_path):
    try:
        # Read the content from the Word document
        doc = Document(docx_path)
        doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return doc_content
    except Exception as e:
        print(f"Error reading document at '{docx_path}': {e}")
        return None

def show_differences(content1, content2):
    # Use difflib to compare the two sets of content
    differ = difflib.Differ()
    diff = list(differ.compare(content1.splitlines(), content2.splitlines()))

    # Collect lines with differences
    webpage_only = [line[2:] for line in diff if line.startswith('+ ')]
    doc_only = [line[2:] for line in diff if line.startswith('- ')]

    return webpage_only, doc_only

def log_differences_to_docx(webpage_only, doc_only, output_path):
    # Create a new Word document
    doc = Document()

    # Add headings for additions and deletions
    doc.add_heading('Things Missing from Webpage', level=1)
    for line in webpage_only:
        doc.add_paragraph(line)

    doc.add_heading('Things Missing from Document', level=1)
    for line in doc_only:
        doc.add_paragraph(line)

    # Save the document
    doc.save(output_path)

def update_csv_with_output_path(csv_path, webpage_url, output_path):
    df = pd.read_csv(csv_path)
    df.loc[df['Webpage URL'] == webpage_url, 'Output file name'] = output_path
    df.to_csv(csv_path, index=False)

def main():
    # Path to the CSV file containing webpage URLs and document paths
    csv_path = r"C:\Users\User1\Downloads\EmploinWebsiteContentQA.csv"

    # Read CSV into DataFrame
    df = pd.read_csv(csv_path)

    for index, row in df.iterrows():
        # Get content from the webpage (text only, without HTML tags)
        webpage_content = get_webpage_content(row['Webpage URL'])

        # Get content from the Word document
        doc_content = get_docx_content(row['Document Path'])

        # Show the differences
        if doc_content is not None:
            webpage_only, doc_only = show_differences(webpage_content, doc_content)
        else:
            # Handle the case where doc_content is None
            print("Skipping comparison due to error in reading document.")

        # Log differences to a new Word document
        output_path = f'differences_output_{index}.docx'
        log_differences_to_docx(webpage_only, doc_only, output_path)

        # Update CSV with the path of the new document
        update_csv_with_output_path(csv_path, row['Webpage URL'], output_path)

if __name__ == "__main__":
    main()
